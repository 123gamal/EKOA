"""File parsers for document ingestion."""

from pathlib import Path

import fitz


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n\n".join(text_parts)


def parse_docx(file_path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    # Include text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n\n".join(paragraphs)


def parse_md(file_path: str) -> str:
    """Read a Markdown file as plain text."""
    return Path(file_path).read_text(encoding="utf-8")


def parse_txt(file_path: str) -> str:
    """Read a plain text file."""
    return Path(file_path).read_text(encoding="utf-8")


def parse_file(file_path: str, content_type: str) -> str:
    """Route to the correct parser based on file content type."""
    ctype = content_type.lower()
    if "pdf" in ctype:
        return parse_pdf(file_path)
    elif file_path.endswith(".docx") or "wordprocessingml" in ctype or "word" in ctype:
        return parse_docx(file_path)
    elif "markdown" in ctype or file_path.endswith(".md"):
        return parse_md(file_path)
    else:
        return parse_txt(file_path)
